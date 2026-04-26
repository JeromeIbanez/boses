import logging

logger = logging.getLogger(__name__)

from fastapi import APIRouter, Depends, File, HTTPException, BackgroundTasks, UploadFile, Request
from sqlalchemy.orm import Session
from sqlalchemy import select

from app.auth.dependencies import CurrentUser, get_current_user
from app.config import settings
from app.database import get_db
from app.services.openai_client import get_openai_client
from app.models.idi_message import IDIMessage
from app.models.persona import Persona
from app.models.simulation import Simulation
from app.models.simulation_briefing import SimulationBriefing
from app.models.simulation_result import SimulationResult
from app.routers.common import get_project_or_404 as _get_project_or_404
from app.limiter import limiter
from app.schemas.simulation import (
    ConjointDesignCreate,
    IDIMessageCreate,
    IDIMessageResponse,
    SimulationCreate,
    SimulationResponse,
    SimulationResultResponse,
)
from app.services.simulation_engine import run_simulation

router = APIRouter(prefix="/projects/{project_id}/simulations", tags=["simulations"])



@router.get("", response_model=list[SimulationResponse])
def list_simulations(
    project_id: str,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    from app.models.reproducibility import ReproducibilityRun
    _get_project_or_404(project_id, db, current_user.company_id)
    reliability_run_ids = select(ReproducibilityRun.simulation_id)
    return db.execute(
        select(Simulation)
        .where(Simulation.project_id == project_id)
        .where(Simulation.id.not_in(reliability_run_ids))
        .order_by(Simulation.created_at.desc())
    ).scalars().all()


@router.post("", response_model=SimulationResponse, status_code=201)
@limiter.limit("20/hour")
def create_simulation(
    request: Request,
    project_id: str,
    body: SimulationCreate,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    _get_project_or_404(project_id, db, current_user.company_id)

    # Validate based on type
    if body.simulation_type == "concept_test":
        if not body.briefing_ids:
            raise HTTPException(status_code=422, detail="At least one briefing is required for concept tests")
        if not body.prompt_question or not body.prompt_question.strip():
            raise HTTPException(status_code=422, detail="prompt_question is required for concept tests")

    if body.simulation_type == "idi_manual":
        if not body.idi_persona_id:
            raise HTTPException(status_code=422, detail="idi_persona_id is required for manual IDI")

    if body.simulation_type == "focus_group":
        if not body.prompt_question or not body.prompt_question.strip():
            raise HTTPException(status_code=422, detail="prompt_question is required for focus groups")

    if body.simulation_type == "conjoint":
        if not body.prompt_question or not body.prompt_question.strip():
            raise HTTPException(status_code=422, detail="prompt_question (product category) is required for conjoint tests")

    # For idi_ai with inline script, validate it's non-empty.
    # File-upload mode sends no script here — the background task is triggered by the upload endpoint instead.
    idi_ai_ready = body.simulation_type == "idi_ai" and bool(body.idi_script_text and body.idi_script_text.strip())

    initial_status = "active" if body.simulation_type == "idi_manual" else "pending"

    # Validate all persona groups exist, belong to this project, and are complete
    from app.models.persona_group import PersonaGroup as PersonaGroupModel
    groups = db.execute(
        select(PersonaGroupModel).where(PersonaGroupModel.id.in_(body.persona_group_ids))
    ).scalars().all()
    if len(groups) != len(body.persona_group_ids):
        raise HTTPException(status_code=422, detail="One or more persona group IDs not found")
    for g in groups:
        if str(g.project_id) != project_id:
            raise HTTPException(status_code=403, detail="Persona group does not belong to this project")
        if g.generation_status != "complete":
            raise HTTPException(status_code=422, detail=f"Persona group '{g.name}' has not finished generating personas yet")

    simulation = Simulation(
        project_id=project_id,
        persona_group_id=body.persona_group_ids[0],  # legacy compat — first group
        prompt_question=body.prompt_question,
        simulation_type=body.simulation_type,
        idi_script_text=body.idi_script_text,
        idi_persona_id=body.idi_persona_id,
        survey_schema=body.survey_schema,
        status=initial_status,
    )
    db.add(simulation)
    db.flush()

    simulation.persona_groups = groups

    if body.briefing_ids:
        from app.models.briefing import Briefing as BriefingModel
        briefings = db.execute(
            select(BriefingModel).where(BriefingModel.id.in_(body.briefing_ids))
        ).scalars().all()
        if len(briefings) != len(body.briefing_ids):
            raise HTTPException(status_code=422, detail="One or more briefing IDs not found")
        for b in briefings:
            if str(b.project_id) != project_id:
                raise HTTPException(status_code=403, detail="Briefing does not belong to this project")
        simulation.briefings = briefings

    db.commit()
    db.refresh(simulation)

    # Schedule background task now only if we already have the script (text mode).
    # File-upload mode (IDI/survey): task is scheduled after the file is processed / confirmed.
    if body.simulation_type in ("concept_test", "focus_group") or idi_ai_ready:
        background_tasks.add_task(run_simulation, simulation_id=str(simulation.id))

    return simulation


@router.get("/convergence")
def get_convergence(
    project_id: str,
    persona_group_id: str,
    briefing_id: str | None = None,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Return convergence scores across completed simulations sharing the same persona group."""
    from app.services.benchmarking_service import compute_convergence
    _get_project_or_404(project_id, db, current_user.company_id)
    return compute_convergence(
        project_id=project_id,
        briefing_id=briefing_id,
        persona_group_id=persona_group_id,
        db=db,
    )


@router.get("/{simulation_id}", response_model=SimulationResponse)
def get_simulation(
    project_id: str,
    simulation_id: str,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    from datetime import datetime, timedelta
    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")

    # If stuck in running/pending for >20 minutes, mark as failed
    if simulation.status in ("running", "pending", "generating_report"):
        age = datetime.utcnow() - simulation.created_at
        if age > timedelta(minutes=20):
            simulation.status = "failed"
            simulation.error_message = "Simulation timed out — the background task may have been interrupted. Please try again."
            db.commit()

    return simulation


@router.post("/{simulation_id}/abort", response_model=SimulationResponse)
def abort_simulation(
    project_id: str,
    simulation_id: str,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Abort a running or pending simulation."""
    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    if simulation.status not in ("pending", "running", "generating_report"):
        raise HTTPException(status_code=422, detail="Simulation is not currently running")
    simulation.status = "failed"
    simulation.error_message = "Aborted by user"
    db.commit()
    db.refresh(simulation)
    return simulation


@router.delete("/{simulation_id}", status_code=204)
def delete_simulation(
    project_id: str,
    simulation_id: str,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    db.delete(simulation)
    db.commit()


@router.get("/{simulation_id}/results", response_model=list[SimulationResultResponse])
def get_simulation_results(
    project_id: str,
    simulation_id: str,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    return db.execute(
        select(SimulationResult)
        .where(SimulationResult.simulation_id == simulation_id)
        .order_by(SimulationResult.created_at)
    ).scalars().all()


@router.post("/{simulation_id}/script", response_model=SimulationResponse)
async def upload_idi_script(
    project_id: str,
    simulation_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Upload a script file (.txt or .docx) for an IDI simulation.
    Extracts the text and stores it in idi_script_text."""
    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    if simulation.simulation_type not in ("idi_ai", "idi_manual"):
        raise HTTPException(status_code=422, detail="Script upload only applies to IDI simulations")

    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in (".txt", ".docx"):
        raise HTTPException(status_code=422, detail="Only .txt and .docx files are supported for scripts")

    # Save file temporarily, extract text
    save_dir = Path(settings.UPLOAD_DIR) / project_id / "scripts"
    save_dir.mkdir(parents=True, exist_ok=True)
    save_path = save_dir / f"{uuid.uuid4().hex}_{filename}"

    with open(save_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    if ext == ".txt":
        raw_text = save_path.read_text(errors="replace")
    else:
        try:
            import docx
            doc = docx.Document(str(save_path))
            raw_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.error(f"Failed to read .docx script: {e}")
            raise HTTPException(status_code=422, detail="Could not read .docx file")

    # Use LLM to extract only the interview questions, regardless of document format
    try:
        client = get_openai_client()
        response = client.chat.completions.create(
            model=settings.OPENAI_MODEL,
            messages=[{
                "role": "user",
                "content": (
                    "Extract only the interview questions from the document below. "
                    "Return one question per line, with no numbering, no interviewer notes, "
                    "no section headings, no instructions, and no other text. "
                    "If something is not a question asked to the respondent, exclude it.\n\n"
                    f"DOCUMENT:\n{raw_text}"
                ),
            }],
            temperature=0,
        )
        script_text = response.choices[0].message.content or raw_text
    except Exception as e:
        logger.error(f"LLM question extraction failed, falling back to raw text: {e}")
        script_text = raw_text

    simulation.idi_script_text = script_text
    db.commit()
    db.refresh(simulation)

    # If this is an idi_ai simulation waiting on its script, kick off the background task now
    if simulation.simulation_type == "idi_ai" and simulation.status == "pending":
        background_tasks.add_task(run_simulation, simulation_id=str(simulation.id))

    return simulation


# ---------------------------------------------------------------------------
# Manual IDI chat endpoints
# ---------------------------------------------------------------------------

@router.get("/{simulation_id}/messages", response_model=list[IDIMessageResponse])
def get_idi_messages(
    project_id: str,
    simulation_id: str,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    return db.execute(
        select(IDIMessage)
        .where(IDIMessage.simulation_id == simulation_id)
        .order_by(IDIMessage.created_at)
    ).scalars().all()


@router.post("/{simulation_id}/messages", response_model=IDIMessageResponse)
def send_idi_message(
    project_id: str,
    simulation_id: str,
    body: IDIMessageCreate,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Send a message as the interviewer and receive the persona's response."""
    from app.services.prompts import idi_system_prompt

    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    if simulation.simulation_type != "idi_manual":
        raise HTTPException(status_code=422, detail="Messages endpoint is only for manual IDI sessions")
    if simulation.status != "active":
        raise HTTPException(status_code=422, detail="This interview session is no longer active")

    persona = db.get(Persona, simulation.idi_persona_id)
    if not persona:
        raise HTTPException(status_code=404, detail="Persona not found")

    # Save user message
    user_msg = IDIMessage(
        simulation_id=simulation_id,
        persona_id=None,
        role="user",
        content=body.content,
    )
    db.add(user_msg)
    db.flush()

    # Build full conversation history for OpenAI
    history = db.execute(
        select(IDIMessage)
        .where(IDIMessage.simulation_id == simulation_id)
        .order_by(IDIMessage.created_at)
    ).scalars().all()

    from app.services.briefing_utils import combine_briefing_texts
    briefing_text = combine_briefing_texts(simulation.briefings)
    system_prompt = idi_system_prompt(persona, briefing_text)

    # Build script context block if available
    script_block = ""
    if simulation.idi_script_text:
        script_block = (
            f"\n\nNote: The interviewer has prepared the following research questions as a guide "
            f"(they may not follow this exactly):\n{simulation.idi_script_text}"
        )

    oai_messages = [{"role": "system", "content": system_prompt + script_block}]
    for msg in history:
        oai_messages.append({
            "role": "user" if msg.role == "user" else "assistant",
            "content": msg.content,
        })

    client = get_openai_client()
    response = client.chat.completions.create(
        model=settings.OPENAI_MODEL,
        messages=oai_messages,
        temperature=0.85,
    )
    answer = response.choices[0].message.content or ""

    # Save persona response
    persona_msg = IDIMessage(
        simulation_id=simulation_id,
        persona_id=persona.id,
        role="persona",
        content=answer,
    )
    db.add(persona_msg)
    db.commit()
    db.refresh(persona_msg)
    return persona_msg


# ---------------------------------------------------------------------------
# Survey endpoints
# ---------------------------------------------------------------------------

@router.post("/{simulation_id}/survey", response_model=SimulationResponse)
async def upload_survey_file(
    project_id: str,
    simulation_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Upload a survey form (.txt or .docx). AI parses it into structured questions.
    Returns the updated simulation with survey_schema populated for preview.
    Does NOT start the simulation — call /run to confirm and start."""
    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    if simulation.simulation_type != "survey":
        raise HTTPException(status_code=422, detail="Survey upload only applies to survey simulations")

    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in (".txt", ".docx"):
        raise HTTPException(status_code=422, detail="Only .txt and .docx files are supported")

    save_dir = Path(settings.UPLOAD_DIR) / project_id / "surveys"
    save_dir.mkdir(parents=True, exist_ok=True)
    save_path = save_dir / f"{uuid.uuid4().hex}_{filename}"

    with open(save_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    if ext == ".txt":
        raw_text = save_path.read_text(errors="replace")
    else:
        try:
            import docx
            doc = docx.Document(str(save_path))
            raw_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.error(f"Failed to read .docx survey: {e}")
            raise HTTPException(status_code=422, detail="Could not read .docx file")

    try:
        import json
        client = get_openai_client()
        response = client.chat.completions.create(
            model=settings.OPENAI_MODEL,
            messages=[{
                "role": "user",
                "content": (
                    "Parse the following survey document into a structured JSON object. "
                    "Return ONLY valid JSON with no markdown, no code blocks, no explanation — just the raw JSON object.\n\n"
                    "The JSON must have this exact shape:\n"
                    '{"questions": [...]}\n\n'
                    "Each question object must have:\n"
                    '- "id": string (q1, q2, q3, ...)\n'
                    '- "type": one of "likert", "multiple_choice", "open_ended"\n'
                    '- "text": the question text\n'
                    "For likert: also include \"scale\" (integer, default 5), \"low_label\" (string), \"high_label\" (string)\n"
                    "For multiple_choice: also include \"options\" (array of strings)\n"
                    "For open_ended: no extra fields needed\n\n"
                    "Infer the type from context: rating/agreement/likelihood scales → likert; "
                    "choose one/select → multiple_choice; free text → open_ended.\n\n"
                    f"SURVEY DOCUMENT:\n{raw_text}"
                ),
            }],
            temperature=0,
        )
        raw_json = response.choices[0].message.content or ""
        survey_schema = json.loads(raw_json)
    except Exception as e:
        logger.error(f"Survey parsing failed: {e}")
        raise HTTPException(status_code=422, detail="Could not parse survey questions from the uploaded file")

    simulation.survey_schema = survey_schema
    db.commit()
    db.refresh(simulation)
    return simulation


@router.post("/{simulation_id}/run", response_model=SimulationResponse)
def run_survey_simulation(
    project_id: str,
    simulation_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Confirm and start a survey simulation after the user has previewed the parsed questions."""
    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    if simulation.simulation_type != "survey":
        raise HTTPException(status_code=422, detail="Run endpoint is only for survey simulations")
    if simulation.status != "pending":
        raise HTTPException(status_code=422, detail="Simulation is not in pending state")
    if not simulation.survey_schema:
        raise HTTPException(status_code=422, detail="Survey questions not uploaded yet")

    questions = (simulation.survey_schema or {}).get("questions", [])
    bad = [q["id"] for q in questions if q.get("type") == "multiple_choice" and not q.get("options")]
    if bad:
        raise HTTPException(
            status_code=422,
            detail=f"Multiple choice question(s) {', '.join(bad)} have no options. Please re-upload the survey or edit the questions."
        )

    background_tasks.add_task(run_simulation, simulation_id=str(simulation.id))
    db.refresh(simulation)
    return simulation


@router.post("/{simulation_id}/conjoint-design", response_model=SimulationResponse)
def run_conjoint_simulation(
    project_id: str,
    simulation_id: str,
    body: ConjointDesignCreate,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Save conjoint attribute/level design and start the simulation."""
    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    if simulation.simulation_type != "conjoint":
        raise HTTPException(status_code=422, detail="This endpoint is only for conjoint simulations")
    if simulation.status != "pending":
        raise HTTPException(status_code=422, detail="Simulation is not in pending state")
    if len(body.attributes) < 2:
        raise HTTPException(status_code=422, detail="At least 2 attributes are required")
    if any(len(a.levels) < 2 for a in body.attributes):
        raise HTTPException(status_code=422, detail="Each attribute must have at least 2 levels")

    simulation.survey_schema = {
        "attributes": [{"name": a.name, "levels": a.levels} for a in body.attributes],
        "n_tasks": min(max(body.n_tasks, 6), 20),
    }
    db.commit()
    db.refresh(simulation)

    background_tasks.add_task(run_simulation, simulation_id=str(simulation.id))
    return simulation


# ---------------------------------------------------------------------------
# Phase 2: Reliability check (reproducibility)
# ---------------------------------------------------------------------------

@router.post("/{simulation_id}/reliability-check")
def create_reliability_check(
    project_id: str,
    simulation_id: str,
    body: dict,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Kick off N repeat runs of a simulation to compute a confidence score."""
    from app.models.reproducibility import ReproducibilityStudy, ReproducibilityRun
    from app.models.simulation import Simulation as SimModel

    _get_project_or_404(project_id, db, current_user.company_id)
    source = db.get(SimModel, simulation_id)
    if not source or str(source.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    if source.status != "complete":
        raise HTTPException(status_code=422, detail="Only completed simulations can be checked for reliability")

    n_runs = max(2, min(int(body.get("n_runs", 3)), 5))

    study = ReproducibilityStudy(
        project_id=project_id,
        source_simulation_id=source.id,
        n_runs=n_runs,
        status="running",
    )
    db.add(study)
    db.flush()

    for i in range(1, n_runs + 1):
        repeat_sim = SimModel(
            project_id=str(source.project_id),
            persona_group_id=source.persona_group_id,
            prompt_question=source.prompt_question,
            simulation_type=source.simulation_type,
            idi_script_text=source.idi_script_text,
            survey_schema=source.survey_schema,
            status="pending",
        )
        db.add(repeat_sim)
        db.flush()
        repeat_sim.persona_groups = list(source.persona_groups)
        repeat_sim.briefings = list(source.briefings)

        repro_run = ReproducibilityRun(
            study_id=study.id,
            simulation_id=repeat_sim.id,
            run_index=i,
        )
        db.add(repro_run)
        background_tasks.add_task(run_simulation, simulation_id=str(repeat_sim.id))

    db.commit()
    db.refresh(study)
    return {
        "id": str(study.id),
        "source_simulation_id": simulation_id,
        "n_runs": n_runs,
        "status": study.status,
        "confidence_score": None,
        "created_at": study.created_at.isoformat(),
    }


@router.get("/{simulation_id}/reliability-check")
def get_reliability_check(
    project_id: str,
    simulation_id: str,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Get the latest reliability study for a simulation."""
    from app.models.reproducibility import ReproducibilityStudy
    from sqlalchemy import select as sa_select

    _get_project_or_404(project_id, db, current_user.company_id)

    study = db.execute(
        sa_select(ReproducibilityStudy)
        .where(ReproducibilityStudy.source_simulation_id == simulation_id)
        .order_by(ReproducibilityStudy.created_at.desc())
        .limit(1)
    ).scalar_one_or_none()

    if not study:
        return {"exists": False}

    return {
        "exists": True,
        "id": str(study.id),
        "source_simulation_id": simulation_id,
        "n_runs": study.n_runs,
        "status": study.status,
        "confidence_score": study.confidence_score,
        "sentiment_agreement_rate": study.sentiment_agreement_rate,
        "distribution_variance_score": study.distribution_variance_score,
        "theme_overlap_coefficient": study.theme_overlap_coefficient,
        "score_breakdown": study.score_breakdown,
        "created_at": study.created_at.isoformat(),
        "completed_at": study.completed_at.isoformat() if study.completed_at else None,
    }


@router.post("/{simulation_id}/end", response_model=SimulationResponse)
def end_idi_session(
    project_id: str,
    simulation_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """End a manual IDI session and trigger report generation."""
    from app.services.idi_engine import generate_idi_report_from_messages

    _get_project_or_404(project_id, db, current_user.company_id)
    simulation = db.get(Simulation, simulation_id)
    if not simulation or str(simulation.project_id) != project_id:
        raise HTTPException(status_code=404, detail="Simulation not found")
    if simulation.simulation_type != "idi_manual":
        raise HTTPException(status_code=422, detail="Only manual IDI sessions can be ended this way")
    if simulation.status != "active":
        raise HTTPException(status_code=422, detail="Session is not active")

    simulation.status = "generating_report"
    db.commit()
    db.refresh(simulation)

    background_tasks.add_task(generate_idi_report_from_messages, simulation_id=simulation_id)
    return simulation
